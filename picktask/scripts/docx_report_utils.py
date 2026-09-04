"""DOCX 报告通用排版工具。"""

from __future__ import annotations

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT_CN = "PingFang SC"
COLOR_TITLE = RGBColor(0x0F, 0x17, 0x2A)
COLOR_H1 = RGBColor(0x1E, 0x3A, 0x5F)
COLOR_H2 = RGBColor(0x2D, 0x4A, 0x6F)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_HEADER_BG = "1E3A5F"
COLOR_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_STRIPE = "F0F4F8"

SZ_TITLE = Pt(26)
SZ_SUBTITLE = Pt(13)
SZ_META = Pt(11)
SZ_H1 = Pt(16)
SZ_H2 = Pt(13)
SZ_BODY = Pt(11)
SZ_TABLE = Pt(10)
SZ_CAPTION = Pt(9.5)
SZ_BULLET = Pt(11)


def set_run_font(run, *, size: Pt, bold: bool = False, color: RGBColor | None = None, name: str = FONT_CN) -> None:
    run.bold = bold
    run.font.size = size
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def fmt_spacing(fmt, *, before: Pt | None = None, after: Pt | None = None, line: float | None = 1.35) -> None:
    if before is not None:
        fmt.space_before = before
    if after is not None:
        fmt.space_after = after
    if line is not None:
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = line


def para_spacing(para, *, before: Pt | None = None, after: Pt | None = None, line: float | None = 1.35) -> None:
    fmt_spacing(para.paragraph_format, before=before, after=after, line=line)


def shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def cell_margins(cell, *, top=80, bottom=80, left=120, right=120) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = SZ_BODY
    normal.font.color.rgb = COLOR_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    fmt_spacing(normal.paragraph_format, before=Pt(0), after=Pt(6), line=1.35)

    for level, size, color, sb, sa in [
        (1, SZ_H1, COLOR_H1, Pt(18), Pt(8)),
        (2, SZ_H2, COLOR_H2, Pt(14), Pt(6)),
        (3, Pt(12), COLOR_H2, Pt(10), Pt(4)),
    ]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = FONT_CN
        h.font.size = size
        h.font.bold = True
        h.font.color.rgb = color
        h._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        fmt_spacing(h.paragraph_format, before=sb, after=sa, line=1.2)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT_CN
    bullet.font.size = SZ_BULLET
    bullet.font.color.rgb = COLOR_BODY
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    fmt_spacing(bullet.paragraph_format, before=Pt(0), after=Pt(3), line=1.3)


def add_body(doc: Document, text: str, *, indent: bool = False) -> None:
    p = doc.add_paragraph()
    para_spacing(p, before=Pt(0), after=Pt(8), line=1.4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=SZ_BODY, color=COLOR_BODY)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    para_spacing(p, before=Pt(2), after=Pt(10), line=1.2)
    run = p.add_run(text)
    set_run_font(run, size=SZ_CAPTION, color=COLOR_MUTED, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        set_run_font(run, size=SZ_BULLET, color=COLOR_BODY)
        para_spacing(p, before=Pt(0), after=Pt(4), line=1.35)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths: list[float] | None = None) -> None:
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths and len(col_widths) == n_cols:
        total_cm = sum(col_widths)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
        table.width = Cm(total_cm)

    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=Pt(2), after=Pt(2), line=1.15)
        run = p.add_run(text)
        set_run_font(run, size=SZ_TABLE, bold=True, color=COLOR_HEADER_FG)
        shade_cell(cell, COLOR_HEADER_BG)
        cell_margins(cell)

    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        stripe = r_idx % 2 == 1
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            para_spacing(p, before=Pt(1), after=Pt(1), line=1.2)
            run = p.add_run(text)
            set_run_font(run, size=SZ_TABLE, color=COLOR_BODY)
            if stripe:
                shade_cell(cell, COLOR_STRIPE)
            cell_margins(cell)

    spacer = doc.add_paragraph()
    para_spacing(spacer, before=Pt(0), after=Pt(12), line=1.0)


def add_cover(doc: Document, *, title: str, subtitle: str, meta_lines: list[str] | None = None) -> None:
    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(title_p, before=Pt(0), after=Pt(16), line=1.0)
    run = title_p.add_run(title)
    set_run_font(run, size=SZ_TITLE, bold=True, color=COLOR_TITLE)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(sub_p, before=Pt(0), after=Pt(24), line=1.4)
    run = sub_p.add_run(subtitle)
    set_run_font(run, size=SZ_SUBTITLE, color=COLOR_H1)

    lines = meta_lines or [f"生成日期：{date.today().strftime('%Y 年 %m 月 %d 日')}"]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=Pt(0), after=Pt(6), line=1.3)
        run = p.add_run(line)
        set_run_font(run, size=SZ_META, color=COLOR_MUTED)

    doc.add_page_break()


def add_dataset_section(
    doc: Document,
    *,
    name: str,
    overview_rows: list[list[str]],
    detail_bullets: list[str],
    strengths: list[str],
    limitations: list[str],
    usage: list[str],
) -> None:
    doc.add_heading(name, level=1)
    add_table(doc, ["属性", "说明"], overview_rows, col_widths=[3.2, 13.2])
    add_caption(doc, "数据集详情")
    add_bullets(doc, detail_bullets)
    add_caption(doc, "优势")
    add_bullets(doc, strengths)
    add_caption(doc, "局限")
    add_bullets(doc, limitations)
    add_caption(doc, "使用建议")
    add_bullets(doc, usage)
