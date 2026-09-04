#!/usr/bin/env python3
"""生成《机器人数据采集现状报告》DOCX（含排版与字号规范）。"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parent
DOCX_DEPS = SCRIPTS_DIR.parent / ".docx_deps"
if DOCX_DEPS.is_dir() and str(DOCX_DEPS) not in sys.path:
    sys.path.insert(0, str(DOCX_DEPS))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

WORKSPACE_DIR = SCRIPTS_DIR.parent.parent
OUTPUT_PATH = WORKSPACE_DIR / "机器人数据采集现状报告.docx"

# ── 字号与颜色规范 ──────────────────────────────────────────
FONT_CN = "PingFang SC"
FONT_EN = "Helvetica Neue"
COLOR_TITLE = RGBColor(0x0F, 0x17, 0x2A)
COLOR_H1 = RGBColor(0x1E, 0x3A, 0x5F)
COLOR_H2 = RGBColor(0x2D, 0x4A, 0x6F)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)
COLOR_HEADER_BG = "1E3A5F"
COLOR_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_STRIPE = "F0F4F8"

SZ_TITLE = Pt(26)
SZ_SUBTITLE = Pt(13)
SZ_META = Pt(11)
SZ_H1 = Pt(16)
SZ_H2 = Pt(13)
SZ_BODY = Pt(11)
SZ_TABLE = Pt(10)
SZ_CAPTION = Pt(9.5)
SZ_BULLET = Pt(11)


def _set_run_font(run, *, size: Pt, bold: bool = False, color: RGBColor | None = None, name: str = FONT_CN) -> None:
    run.bold = bold
    run.font.size = size
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def _fmt_spacing(fmt, *, before: Pt | None = None, after: Pt | None = None, line: float | None = 1.35) -> None:
    if before is not None:
        fmt.space_before = before
    if after is not None:
        fmt.space_after = after
    if line is not None:
        fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        fmt.line_spacing = line


def _para_spacing(para, *, before: Pt | None = None, after: Pt | None = None, line: float | None = 1.35) -> None:
    _fmt_spacing(para.paragraph_format, before=before, after=after, line=line)


def _shade_cell(cell, fill_hex: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _cell_margins(cell, *, top=80, bottom=80, left=120, right=120) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = SZ_BODY
    normal.font.color.rgb = COLOR_BODY
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    _fmt_spacing(normal.paragraph_format, before=Pt(0), after=Pt(6), line=1.35)

    for level, size, color, sb, sa in [
        (1, SZ_H1, COLOR_H1, Pt(18), Pt(8)),
        (2, SZ_H2, COLOR_H2, Pt(14), Pt(6)),
        (3, Pt(12), COLOR_H2, Pt(10), Pt(4)),
    ]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = FONT_CN
        h.font.size = size
        h.font.bold = True
        h.font.color.rgb = color
        h._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        _fmt_spacing(h.paragraph_format, before=sb, after=sa, line=1.2)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT_CN
    bullet.font.size = SZ_BULLET
    bullet.font.color.rgb = COLOR_BODY
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    _fmt_spacing(bullet.paragraph_format, before=Pt(0), after=Pt(3), line=1.3)


def add_body(doc: Document, text: str, *, indent: bool = False) -> None:
    p = doc.add_paragraph()
    _para_spacing(p, before=Pt(0), after=Pt(8), line=1.4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    _set_run_font(run, size=SZ_BODY, color=COLOR_BODY)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_spacing(p, before=Pt(2), after=Pt(10), line=1.2)
    run = p.add_run(text)
    _set_run_font(run, size=SZ_CAPTION, color=COLOR_MUTED, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        run = p.add_run(item)
        _set_run_font(run, size=SZ_BULLET, color=COLOR_BODY)
        _para_spacing(p, before=Pt(0), after=Pt(4), line=1.35)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths: list[float] | None = None) -> None:
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths and len(col_widths) == n_cols:
        total_cm = sum(col_widths)
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
        table.width = Cm(total_cm)

    # 表头
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p, before=Pt(2), after=Pt(2), line=1.15)
        run = p.add_run(text)
        _set_run_font(run, size=SZ_TABLE, bold=True, color=COLOR_HEADER_FG)
        _shade_cell(cell, COLOR_HEADER_BG)
        _cell_margins(cell)

    # 数据行
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        stripe = r_idx % 2 == 1
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _para_spacing(p, before=Pt(1), after=Pt(1), line=1.2)
            run = p.add_run(text)
            _set_run_font(run, size=SZ_TABLE, color=COLOR_BODY)
            if stripe:
                _shade_cell(cell, COLOR_STRIPE)
            _cell_margins(cell)

    spacer = doc.add_paragraph()
    _para_spacing(spacer, before=Pt(0), after=Pt(12), line=1.0)


def add_cover(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(title_p, before=Pt(0), after=Pt(16), line=1.0)
    run = title_p.add_run("机器人数据采集现状报告")
    _set_run_font(run, size=SZ_TITLE, bold=True, color=COLOR_TITLE)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(sub_p, before=Pt(0), after=Pt(24), line=1.4)
    run = sub_p.add_run("具身智能 / Physical AI 数据生态\n硬件选型 · 格式对比 · 外包评估")
    _set_run_font(run, size=SZ_SUBTITLE, color=COLOR_H1)

    meta_lines = [
        f"生成日期：{date.today().strftime('%Y 年 %m 月 %d 日')}",
        "项目上下文：MuJoCo pickcup 抓杯任务",
        "数据格式：LeRobot v3 · ACT · Residual RL",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p, before=Pt(0), after=Pt(6), line=1.3)
        run = p.add_run(line)
        _set_run_font(run, size=SZ_META, color=COLOR_MUTED)

    doc.add_page_break()


def build_report() -> Document:
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    add_cover(doc)

    doc.add_heading("摘要", level=1)
    add_body(
        doc,
        "本报告梳理 2025–2026 年机器人（具身智能）数据采集行业现状，涵盖采集方式、主要商业公司与开源生态，"
        "并补充 teleop 硬件选型指南、AgiBot World 与 DROID 数据格式差异，以及外包数据采集服务的报价评估框架。"
        "报告结合 MuJoCo pickcup 项目（scripted 示教 → LeRobot → ACT → residual RL）说明仿真采集在产业链中的位置。",
        indent=True,
    )

    doc.add_heading("1. 为什么数据是具身智能的核心瓶颈", level=1)
    add_bullets(doc, [
        "采集成本高：遥操作站单套约 5 万–15 万美元，需熟练操作员、场地与机器人折旧。",
        "采集效率低：单 teleop 工位常见 <200 demonstrations/天；百万级数据需工厂化运营。",
        "质量难控：动作合理性、传感器时间对齐、成功/失败标注、长尾场景覆盖。",
        "跨平台难复用：不同 embodiment、相机布局、控制接口不统一。",
        "Sim-to-Real 存在 gap：仿真可规模化，但接触动力学、视觉域与真机仍有差异。",
    ])
    add_body(
        doc,
        "行业共识：大模型时代的「燃料」在具身领域极度稀缺；数据工厂、外包服务与开源百万级数据集成为 2024–2026 年主赛道。",
    )

    doc.add_heading("2. 主流数据采集方式", level=1)
    add_table(
        doc,
        ["方式", "内容", "优点", "缺点", "代表"],
        [
            ["Robot Teleop", "人直接控真机，记录 state-action", "与部署 embodiment 一致", "贵、慢、难 scale", "DROID、AgiBot"],
            ["Human Demo", "第一人称视频 / UMI / 手套", "场景多样、成本低", "需跨 embodiment 迁移", "Ego4D、Claru"],
            ["Simulation", "MuJoCo / Isaac 规则或状态机", "便宜、可严格判成功", "与真机有 gap", "pickcup scripted"],
            ["VR Teleop", "VR 控仿真或真机", "灵活、部署快", "动力学可能不一致", "Roborax 等"],
            ["HITL / RLHF", "机器人试 + 人标注纠偏", "适合 RL、长尾", "运营复杂", "Scale 等"],
        ],
        col_widths=[2.4, 3.6, 2.8, 2.8, 2.8],
    )
    add_body(doc, "行业常采用「三层金字塔」数据策略：")
    add_bullets(doc, [
        "顶层：目标机器人 in-domain teleop（最贵、最有用）。",
        "中层：通用臂 / 异构机 teleop（如 GELLO、Franka）。",
        "底层：人类 egocentric / 可穿戴（规模最大，需迁移学习）。",
    ])

    doc.add_heading("3. 主要数据采集公司与机构", level=1)

    doc.add_heading("3.1 国际数据服务公司（外包 / 平台）", level=2)
    add_table(
        doc,
        ["公司", "定位", "特点"],
        [
            ["Scale AI", "Physical AI 数据引擎", "全球采集网络、teleop 站、petabyte 级 ingest"],
            ["XDOF", "数据 pipeline + 工具", "2024 成立；GELLO 类 teleop；开源 ABC 数据集"],
            ["Roborax", "BPO 式交付", "12 国 41 中心；专用 / 众包 / 混合；4 周上线"],
            ["Human Data Co.", "通用人形数据", "UMI、teleop、LeRobot 兼容；HITL 标注"],
            ["Humanola", "Teleop + API 平台", "低延迟遥操、多模态处理"],
            ["Proxy Robotics", "Teleop 基础设施", "跨洲低延迟；数据进 ML pipeline"],
            ["Claru", "Human demo / 视频", "1 万+ 采集员；egocentric 多场景"],
            ["Human Archive", "多模态对齐", "YC；定制硬件 + 大规模标注"],
        ],
        col_widths=[3.0, 3.5, 9.9],
    )

    doc.add_heading("3.2 中国：本体厂 + 数据工厂", level=2)
    add_bullets(doc, [
        "智元 AgiBot：上海张江 ~2000–3000㎡ 采集工厂；开源 AgiBot World（百万级真机轨迹、850TB+）。",
        "国地中心 / 各地人形机器人创新中心：大规模训练场与采集基地。",
        "傅利叶、宇树等：本体 + 遥操能力；行业讨论重点在数据质量评判而非仅数量。",
        "上海 AI Lab、库帕思等：与智元联合做数据与平台。",
    ])

    doc.add_heading("3.3 开源数据集（生态基础设施）", level=2)
    add_table(
        doc,
        ["项目", "规模", "说明"],
        [
            ["Open X-Embodiment", "100 万+ episodes，22+ 机型", "Google DeepMind + 33 机构；RT-X"],
            ["DROID", "7.6 万轨迹，564 场景", "13 机构统一 Franka + Quest 2 teleop"],
            ["BridgeData V2", "~6 万", "Berkeley / Stanford，WidowX"],
            ["AgiBot World", "百万级（目标）", "智元人形机；家居 / 工业 / 商超等"],
        ],
        col_widths=[3.5, 4.5, 8.4],
    )

    doc.add_heading("3.4 行业趋势（2025–2026）", level=2)
    add_bullets(doc, [
        "从实验室小批量 → 工厂化 BPO（Roborax、Scale、XDOF）。",
        "Teleop 仍是 gold standard；egocentric 视频用于预训练 + 少量 in-domain fine-tune。",
        "质量 > 数量：对抗式采集、长尾补采、多轮 QA 成为差异化能力。",
        "格式标准化加速：LeRobot v2/v3、RLDS、HDF5；交付需直接进训练 pipeline。",
        "中国路径：物理数据工厂 + 政府推动 + 开源抢生态；美国路径：startup 卖 data-as-a-service。",
    ])

    doc.add_heading("4. 与 MuJoCo pickcup 项目的对照", level=1)
    add_table(
        doc,
        ["维度", "pickcup 项目", "工业级采集"],
        [
            ["采集方式", "MuJoCo scripted 状态机", "真机 teleop 工厂"],
            ["规模", "200 条成功 episode", "10⁴–10⁶ 条"],
            ["观测", "8 维 proprio + head_camera", "多相机 + 力 / 触觉等"],
            ["成本", "电费 + CPU 训练时间", "场地 + 机器人 + 人力"],
            ["目标", "跑通 IL/RL 链路、算法验证", "训 VLA / 通用人形"],
        ],
        col_widths=[3.0, 6.5, 6.9],
    )
    add_body(
        doc,
        "pickcup 使用 LeRobot v3：observation.images.head_camera（640×480, 30fps）、"
        "observation.state / action 均为 8 维（腰 + 右臂 + 夹爪，不含杯位姿）。"
        "这是产业链中最轻量、最适合原型验证的一层。",
    )

    doc.add_page_break()

    doc.add_heading("5. Teleop 硬件选型指南", level=1)
    add_body(
        doc,
        "选型核心问题：你需要的是「与目标机器人 action 空间对齐的真机轨迹」，还是「大规模人类行为视频」？"
        "前者选 leader-follower 或同构 teleop；后者选 UMI / egocentric 可穿戴。",
    )

    doc.add_heading("5.1 主流方案对比", level=2)
    add_table(
        doc,
        ["方案", "原理", "典型成本", "数据质量", "适用场景"],
        [
            ["ALOHA / ALOHA 2", "双 WidowX leader → ViperX follower", "~$20k–30k", "高；双臂精细", "固定工位 bimanual IL"],
            ["GELLO", "3D 打印同构控制器", "~$1k–5k", "中高；单臂", "实验室 teleop，XDOF 生态"],
            ["UMI", "手持夹爪 + 相机", "~$数百–数千", "中；需迁移", "in-the-wild 人类 demo"],
            ["Quest 2 + Franka", "VR 手柄控臂（DROID）", "~$50k+", "高；单臂", "分布式多场景采集"],
            ["VR / 外骨骼", "非同构映射", "视方案", "中", "人形机、远程 teleop"],
            ["键盘 / 关节增量", "软件 teleop", "几乎为 0", "低–中", "仿真调试（pickcup）"],
        ],
        col_widths=[2.6, 3.8, 2.0, 2.4, 4.6],
    )

    doc.add_heading("5.2 选型决策树", level=2)
    add_bullets(doc, [
        "目标机器人已确定 + 需要直接 BC/VLA 训练 → leader-follower（ALOHA/GELLO）或 OEM teleop。",
        "需要快速覆盖多场景、预算有限 → UMI / egocentric + 迁移学习；后续补少量 in-domain teleop。",
        "双臂精细操作（折叠、插入）→ ALOHA 2 类 bimanual 工位。",
        "分布式、多地理场景 → DROID 式标准化硬件（Franka + 相机 + Quest）。",
        "人形 / 全尺寸 → VR teleop 或外骨骼；通常外包 Roborax / Proxy / Scale。",
        "仅算法验证 / 无真机 → MuJoCo scripted（pickcup）或 ALOHA 2 仿真采集。",
    ])

    doc.add_heading("5.3 硬件选型检查清单", level=2)
    add_bullets(doc, [
        "Action 空间是否与策略输出一致（关节角 vs EEF pose vs delta）？",
        "相机数量、分辨率、内外参标定是否满足 VLA 需求？",
        "控制频率（通常 10–30 Hz）与时间戳对齐方案？",
        "数据导出格式：HDF5 / RLDS / LeRobot v2/v3？",
        "成功 / 失败判定：人工 discard 还是自动检测？",
        "工位吞吐量：目标 episodes/天/站？",
        "维护与备件：长时间采集的可靠性。",
    ])

    doc.add_heading("6. AgiBot World vs DROID 格式差异", level=1)

    doc.add_heading("6.1 数据集概览", level=2)
    add_table(
        doc,
        ["维度", "DROID", "AgiBot World"],
        [
            ["主导机构", "13 北美研究机构", "智元 + 上海 AI Lab + 国地中心"],
            ["机器人", "统一 Franka Panda 7DoF", "智元人形机（多 DOF）"],
            ["规模", "~76k 轨迹，564 场景", "百万级轨迹（分批发布）"],
            ["采集方式", "Quest 2 VR teleop", "手柄遥操 + 工厂 QA"],
            ["任务特点", "桌面 manipulation", "长程任务 60–150s，80+ 技能"],
            ["开源许可", "CC-BY 4.0", "研究开源（见各子集）"],
        ],
        col_widths=[3.0, 5.5, 7.9],
    )

    doc.add_heading("6.2 原生存储格式", level=2)
    add_caption(doc, "DROID（原生）")
    add_bullets(doc, [
        "RLDS / TFRecord 为主；LeRobot 提供 droid_1.0.1（v3 parquet + mp4）转换版。",
        "相机：2× Zed 2 外部 + 1× Zed Mini 腕部；控制 10–15 Hz。",
        "观测：关节状态 + 多路 RGB；动作：相对/绝对 EEF 或 joint。",
        "语言：每 episode 多条 instruction 变体。",
        "特点：跨 13 站点硬件完全一致，场景多样性是核心卖点。",
    ])
    add_caption(doc, "AgiBot World（原生）")
    add_bullets(doc, [
        "目录：task_id / episode_id / proprio_stats.h5 + task_*.json。",
        "H5 内嵌丰富 proprio：joint、effector、EEF pose/wrench、head、waist 等。",
        "JSON 含 skill 分段（Pick/Place）、object、帧级 keyframe。",
        "Alpha/Beta 需脚本转 LeRobot；AgiBotWorld2026 已提供 v2.1 结构。",
    ])

    doc.add_heading("6.3 LeRobot 格式对照", level=2)
    add_table(
        doc,
        ["字段", "DROID", "AgiBot", "pickcup"],
        [
            ["robot_type", "franka / OXE tag", "agibot", "unitree_g1_23dof"],
            ["图像键", "exterior_*, wrist_*", "observation.images.*", "head_camera"],
            ["state 维度", "关节 / EEF", "高维（腰+臂+头+手）", "8 维 proprio"],
            ["action", "joint 或 EEF delta", "全关节目标", "8 维关节目标"],
            ["fps", "通常 15", "30", "30"],
            ["版本", "v2/v3", "v2.0–v2.1 / 需转 v3", "v3"],
        ],
        col_widths=[2.4, 3.6, 3.6, 3.6],
    )

    doc.add_heading("6.4 使用建议", level=2)
    add_bullets(doc, [
        "训练 ACT/π0/GR00T：确认 embodiment tag 与数据集一致，不可混用 action 语义。",
        "AgiBot 原生 H5：需 embodied-data 或官方脚本转 LeRobot。",
        "DROID：优先用 LeRobot droid_1.0.1；注意 video key 与 model config 映射。",
        "迁移到 G1 pickcup：不能零样本部署；仅作 pretrain 或架构参考。",
    ])

    doc.add_page_break()

    doc.add_heading("7. 外包数据采集报价评估框架", level=1)
    add_body(
        doc,
        "外包报价差异巨大（从数千美元 pilot 到百万美元年度合同）。"
        "评估时应拆分为：数据采集、标注、格式转换、IP 归属、SLA、复采与 exclusivity。",
    )

    doc.add_heading("7.1 报价模型常见类型", level=2)
    add_table(
        doc,
        ["计费方式", "典型区间（参考）", "适用"],
        [
            ["按 demonstration 条数", "$10–$80 / 条", "teleop 真机"],
            ["按采集小时", "$80–$300 / 小时", "BPO 模式"],
            ["按 episode + QA", "含 discard 率溢价", "高质量工业场景"],
            ["Pilot 包", "$2.5k–$20k / 200–500 条", "验证供应商"],
            ["年度框架", "按 TB 或百万帧", "大规模 pretrain"],
        ],
        col_widths=[4.0, 4.5, 7.9],
    )
    add_caption(doc, "注：以上为 2025–2026 公开市场参考区间，实际需 RFP 询价。")

    doc.add_heading("7.2 评估维度（RFI/RFP 检查清单）", level=2)
    add_bullets(doc, [
        "Embodiment 匹配：是否在目标机器人上采集？",
        "场景覆盖：场景数、物体数、长尾任务比例。",
        "吞吐量：每工位 episodes/天；上线时间线。",
        "QA 流程：discard 率、审核标准、时间对齐校验。",
        "交付格式：LeRobot v2/v3、RLDS、HDF5；是否含 stats.json。",
        "IP 与 exclusivity：数据是否独家？能否开源 / 发论文？",
        "合规：GDPR、数据出境、生物特征脱敏。",
    ])

    doc.add_heading("7.3 性价比计算建议", level=2)
    add_body(doc, "有效成本 = 总报价 ÷ (合格 episodes × 平均有效帧数 × 任务覆盖系数)", indent=True)
    add_bullets(doc, [
        "任务覆盖系数：单场景 0.3–0.5；多场景长尾齐全接近 1.0。",
        "Break-even：<5000 条一次性需求，外包 pilot 通常优于自建 ALOHA。",
        ">100k 条且 embodiment 固定：自建工厂 + 外包 QA 混合更常见。",
    ])

    doc.add_heading("7.4 Red Flags（避坑）", level=2)
    add_bullets(doc, [
        "仅承诺「小时数」不承诺合格条数或 success 定义。",
        "无法提供 sample episode 做格式验证。",
        "action/observation 语义文档缺失。",
        "多相机时间戳不同步。",
        "只做 egocentric 视频却按 teleop 价格报价。",
    ])

    doc.add_heading("7.5 推荐采购路径", level=2)
    add_table(
        doc,
        ["阶段", "建议", "预算量级"],
        [
            ["算法验证", "仿真 scripted + 开源小集", "$0–1k"],
            ["PoC", "外包 pilot 200–500 条 或 GELLO 自建", "$2k–30k"],
            ["Product baseline", "1–5 工位 teleop 或 hybrid 外包", "$50k–500k"],
            ["Foundation pretrain", "Scale/XDOF 框架 + OXE/AgiBot", "$1M+"],
        ],
        col_widths=[3.0, 7.5, 5.9],
    )

    doc.add_heading("8. 结论与建议", level=1)
    add_bullets(doc, [
        "数据采集正从 lab 自给自足变为工厂化、平台化。",
        "Teleop 硬件选型取决于 embodiment 对齐 vs 规模需求。",
        "AgiBot 与 DROID 格式差异大，训练前必须统一 schema。",
        "外包评估应基于合格 episode 有效成本。",
        "pickcup：ACT eval baseline 后再决定是否真机补数据或 residual RL。",
    ])

    doc.add_heading("参考资料", level=1)
    add_bullets(doc, [
        "DROID: https://droid-dataset.github.io/",
        "Open X-Embodiment: https://robotics-transformer-x.github.io/",
        "AgiBot World: https://huggingface.co/agibot-world",
        "Scale Physical AI: https://scale.com/physical-ai",
        "Roborax: https://www.roborax.ai/",
        "ALOHA 2: https://aloha-2.github.io/",
        "LeRobot: https://github.com/huggingface/lerobot",
        "本项目 picktask/readme.md",
    ])

    return doc


def main() -> None:
    doc = build_report()
    doc.save(OUTPUT_PATH)
    print(f"已生成报告: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
