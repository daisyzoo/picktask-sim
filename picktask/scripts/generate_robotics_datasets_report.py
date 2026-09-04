#!/usr/bin/env python3
"""生成《机器人领域重要数据集调研报告》DOCX。"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

SCRIPTS_DIR = Path(__file__).resolve().parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from docx import Document

from docx_report_utils import (
    add_body,
    add_bullets,
    add_cover,
    add_caption,
    add_dataset_section,
    add_table,
    setup_page,
    setup_styles,
)

WORKSPACE_DIR = SCRIPTS_DIR.parent.parent
OUTPUT_PATH = WORKSPACE_DIR / "机器人领域重要数据集调研报告.docx"


def build_report() -> Document:
    doc = Document()
    setup_page(doc)
    setup_styles(doc)
    add_cover(
        doc,
        title="机器人领域重要数据集调研报告",
        subtitle="真机 Manipulation · 跨 Embodiment · 人形机 · 仿真 Benchmark\n格式对比 · 选型指南 · 使用建议",
        meta_lines=[
            f"生成日期：{date.today().strftime('%Y 年 %m 月 %d 日')}",
            "适用读者：具身智能 / VLA / 模仿学习研发",
            "项目参考：MuJoCo pickcup（LeRobot v3 + ACT）",
        ],
    )

    doc.add_heading("摘要", level=1)
    add_body(
        doc,
        "本报告系统介绍 2024–2026 年机器人操作（manipulation）领域最重要的开源与半开源数据集，"
        "涵盖 Open X-Embodiment、DROID、AgiBot World、BridgeData V2、RH20T、RoboMIND、LIBERO 等，"
        "并从规模、embodiment、模态、存储格式、许可证与训练 pipeline 兼容性给出使用建议。"
        "报告末尾提供按研究/产品目标的选型矩阵，并对 MuJoCo pickcup 自采小数据集在生态中的定位作出说明。",
        indent=True,
    )

    doc.add_heading("1. 数据集生态总览", level=1)
    add_body(
        doc,
        "机器人数据集可按「多样性轴」分类：跨机器人（embodiment）、跨场景（scene）、跨任务（task）、"
        "跨模态（RGB / 深度 / 力触觉 / 语言）。没有单一「最好」的数据集，选型取决于你的泛化目标。",
    )
    add_table(
        doc,
        ["数据集", "规模", "Embodiment", "核心卖点", "LeRobot"],
        [
            ["Open X-Embodiment", "~1M+ traj", "22+ 机型", "跨机 pretrain", "IPEC 社区转换"],
            ["DROID", "76k traj", "Franka 统一", "564 真实场景", "droid_1.0.1 v3"],
            ["AgiBot World", "1M+ traj", "智元人形", "长程 + 工厂 QA", "需转换 v2/v3"],
            ["RoboMIND", "107k traj", "4 机型", "失败轨迹 + 细粒度语言", "any4lerobot"],
            ["BridgeData V2", "~60k", "WidowX", "BC baseline 经典", "社区转换"],
            ["RH20T", "110k seq", "Franka 等", "力 / 触觉多模态", "HDF5 原生"],
            ["LIBERO", "仿真", "Franka sim", "终身 / 泛化 benchmark", "LIBERO2lerobot"],
            ["ARIO", "~3M", "双臂等", "中国大规模聚合", "部分 LeRobot"],
            ["pickcup (本地)", "200 ep", "G1 仿真", "domain IL 验证", "原生 v3"],
        ],
        col_widths=[3.2, 2.0, 2.4, 3.8, 3.0],
    )

    doc.add_heading("2. 各数据集详细介绍", level=1)

    add_dataset_section(
        doc,
        name="2.1 Open X-Embodiment (OXE)",
        overview_rows=[
            ["发布方", "Google DeepMind + 21/34 研究机构（2023）"],
            ["规模", "100 万+ 真实机器人轨迹；527 skills；60+ 子数据集聚合"],
            ["机器人", "22 种：Franka、xArm、WidowX、ALOHA 双臂、四足等"],
            ["格式", "RLDS（TensorFlow）；LeRobot 社区有 IPEC-COMMUNITY/OpenX 转换"],
            ["许可", "各子集 CC-BY / Apache 2.0 等（需逐集确认）"],
            ["访问", "robotics-transformer-x.github.io；HuggingFace IPEC 集合"],
        ],
        detail_bullets=[
            "OXE 是首个大规模「跨 embodiment 聚合」标准，动机类似 NLP/CV 的预训练语料 consolidation。",
            "训练 RT-1-X / RT-2-X 等 generalist policy 的标准 pretrain 混合（robot data mixture）。",
            "场景与行为覆盖广，但单个子集质量、相机、action 语义差异大，需按 mixture weight 筛选。",
            "适合研究：cross-embodiment transfer、scaling law、VLA 预训练。",
        ],
        strengths=[
            "Embodiment 与任务多样性业界最高，是 VLA 论文最常引用的 pretrain 基准。",
            "与 RT-X 模型线深度绑定，复现路径清晰。",
            "子集可按 robot type / task 过滤，支持 curriculum 或 stratified sampling。",
        ],
        limitations=[
            "异构严重：action 空间不统一（joint / EEF / delta 混用），训练前必须 canonicalize。",
            "单场景深度不如 DROID；contact-rich 传感不如 RH20T。",
            "全量下载与预处理成本高（存储 + 算力）。",
        ],
        usage=[
            "目标：训 generalist VLA / RT 系模型 → 作 pretrain，再在目标任务 fine-tune。",
            "使用 LeRobot：优先 IPEC-COMMUNITY/openx-lerobot；或 openx2lerobot 工具链。",
            "勿与 pickcup 8 维 G1 数据直接混训；embodiment tag 与 action 语义完全不同。",
            "Fine-tune 前建议：先用与目标机形态接近的子集（如 Franka 类）做 ablation。",
        ],
    )

    add_dataset_section(
        doc,
        name="2.2 DROID",
        overview_rows=[
            ["发布方", "Stanford 等 13 北美机构（2024, RSS）"],
            ["规模", "76,000 轨迹；350 小时；86 tasks；564 scenes；~1.7TB"],
            ["机器人", "统一 Franka Panda 7DoF + 标准相机 + Quest 2 VR teleop"],
            ["格式", "原生 RLDS；LeRobot lerobot/droid_1.0.1（v3 parquet + mp4）"],
            ["许可", "CC-BY 4.0"],
            ["访问", "droid-dataset.github.io；HuggingFace datasets"],
        ],
        detail_bullets=[
            "强调 in-the-wild：52 栋建筑、北美/亚洲/欧洲 50 名采集员、12 个月分布式采集。",
            "硬件完全一致 → 场景多样性是变量，适合研究 scene generalization。",
            "相机：2× Zed 2 外部 stereo + 1× Zed Mini 腕部；控制 10–15 Hz。",
            "语言：每 episode 多条 instruction 变体；支持 language-conditioned policy。",
            "π0.5-DROID、NVIDIA GR00T-N1.7-DROID 等 SOTA 模型的关键 fine-tune 集。",
        ],
        strengths=[
            "真实场景多样性在单臂 manipulation 数据集中顶尖。",
            "标准化硬件使跨 lab 复现与对比公平。",
            "LeRobot v3 官方支持较好；与 GR00T / OpenPI 示例齐全。",
        ],
        limitations=[
            "仅 Franka 单臂；无法直接迁移到人形 G1 或双臂。",
            "无 force/tactile；contact-rich 任务信息主要靠视觉隐式学习。",
            "全量 358GB+，本地训练需足够存储与 video decode 能力。",
        ],
        usage=[
            "目标：单臂 tabletop / 室内 manipulation 泛化 → pretrain 或 co-train 首选之一。",
            "ACT / Diffusion：用 LeRobot 加载，注意 OXE_DROID embodiment tag 与 action 相对/绝对配置。",
            "Eval：RoboArena benchmark 与 π0.5-DROID 权重可作为 generalist 对照。",
            "pickcup 项目：可作 architecture 与 training recipe 参考，不可直接混数据。",
        ],
    )

    add_dataset_section(
        doc,
        name="2.3 AgiBot World",
        overview_rows=[
            ["发布方", "智元机器人 + 上海 AI Lab + 国地中心 + 库帕思（2024.12）"],
            ["规模", "百万级真机轨迹；850TB+；100+ 场景类型；80+ 技能"],
            ["机器人", "智元人形机（腰/头/双臂/灵巧手）"],
            ["格式", "原生 HDF5 (proprio_stats.h5) + JSON；LeRobot v2.1/v3 需转换"],
            ["许可", "研究开源（Alpha/Beta/2026 子集条款略有不同）"],
            ["访问", "huggingface.co/agibot-world；agibot-world.com"],
        ],
        detail_bullets=[
            "中国首个工业级百万真机开源集，上海张江数据工厂采集，多轮人工 QA。",
            "场景：家居 40%、餐饮/工业各 20%、商超/办公各 10%；含 3000+ 真实物体。",
            "长程任务为主：80% 轨迹 60–150 秒；skill 分段标注（Pick/Place 等）。",
            "H5 内含 joint、effector、EEF wrench、head、waist、robot base 等丰富 proprio。",
            "AgiBotWorld2026 子集已按 LeRobot v2.1 提供 meta/info.json + parquet + mp4。",
        ],
        strengths=[
            "规模与长程复杂度领先，适合研究 data scaling law on humanoid。",
            "场景贴近中国家庭/工业/零售，本土部署 relevance 高。",
            "标注深度高：skill 段、object、keyframe，利于 hierarchical VLA。",
        ],
        limitations=[
            "原生格式与 LeRobot v3 / ACT 不完全即开即用，转换链路需维护。",
            "主要为智元 embodiment；迁移到其他 humanoid 仍需大量 fine-tune。",
            "下载与存储门槛极高；Beta 分批发布需关注版本一致性。",
        ],
        usage=[
            "目标：人形机 long-horizon manipulation、中国场景 VLA → pretrain 或 benchmark。",
            "转换：any4lerobot/agibot2lerobot 或 embodied-data convert --from agibot --to lerobot-v3。",
            "与 OXE 对比：AgiBot 单组织质量更统一；OXE 跨 institution 更广。",
            "pickcup：若未来 G1 真机 teleop，可参考 AgiBot 的 QA 与场景划分，而非直接混训。",
        ],
    )

    doc.add_page_break()

    add_dataset_section(
        doc,
        name="2.4 RoboMIND",
        overview_rows=[
            ["发布方", "人形机器人创新中心等（RSS 2025）"],
            ["规模", "107k 成功轨迹 + 5k 失败轨迹；479 tasks；96 object classes"],
            ["机器人", "Franka、UR5e、AgileX 双臂、天工人形"],
            ["格式", "原生自定义；Traly/RoboMIND-lerobot 社区 v3 转换（部分）"],
            ["许可", "见 HuggingFace 官方 repo"],
            ["访问", "huggingface.co/datasets/x-humanoid-robomind/RoboMIND"],
        ],
        detail_bullets=[
            "强调「Normative Data」：统一采集规范，便于 cross-embodiment benchmark。",
            "含 1 万帧级细粒度语言标注；支持 hierarchical language conditioning 研究。",
            "公开失败轨迹 → 类似 RLHF / 负样本学习，在同类数据集中较独特。",
            "任务分 6 大类：关节操作、协调、基础抓取、多物体、精细、场景理解。",
        ],
        strengths=[
            "多 embodiment 在同一规范下，适合 fair comparison。",
            "失败 + 成功成对发布，利于 robust policy 与 reward learning。",
            "语言标注深度超过 DROID / Bridge。",
        ],
        limitations=[
            "LeRobot 全量转换尚未官方完成；社区子集覆盖 4 种 embodiment。",
            "总规模小于 OXE / AgiBot；单任务样本深度因 task 多而较薄。",
        ],
        usage=[
            "目标：multi-embodiment benchmark、language hierarchy、failure-aware IL。",
            "转换：any4lerobot/robomind2lerobot；训练前确认 joint vs EEF action。",
            "可与 OXE 互补：RoboMIND 规范强；OXE 规模大。",
        ],
    )

    add_dataset_section(
        doc,
        name="2.5 BridgeData V2",
        overview_rows=[
            ["发布方", "UC Berkeley / Stanford 等（2023）"],
            ["规模", "~60,000 demonstrations；24 场景；13 skills"],
            ["机器人", "WidowX 6DoF 低成本臂"],
            ["格式", "原生 TFRecord / 自定义；LeRobot 社区有转换"],
            ["许可", "CC-BY 4.0"],
            ["访问", "rail.eecs.berkeley.edu；HuggingFace"],
        ],
        detail_bullets=[
            "Berkeley 系 BC（Behavior Cloning）文化的奠基数据集之一。",
            " tabletop 操作为主，language-conditioned，场景相对固定。",
            "RT-1、Octo 等早期 generalist 的重要组成子集（也并入 OXE）。",
        ],
        strengths=[
            "文档与 baseline 成熟；适合入门 imitation learning 与 reproducing BC 论文。",
            "任务定义清晰，训练/debug 周期短。",
        ],
        limitations=[
            "场景与 embodiment 多样性有限；不适合声称 strong real-world generalization。",
            "已被更大数据集部分替代，但作为 ablation 仍常用。",
        ],
        usage=[
            "目标：快速验证 BC/ACT/Diffusion pipeline → smoke 训练首选之一。",
            "教学与算法课设：数据小、下载快、GPU 友好。",
            "产品级 generalist：应作辅助，不应作为唯一 pretrain。",
        ],
    )

    add_dataset_section(
        doc,
        name="2.6 RH20T",
        overview_rows=[
            ["发布方", "上海交通大学等（2023）"],
            ["规模", "110,000+ 序列；20 tasks；7 场景类型；140 skills"],
            ["机器人", "Franka 等；强调 dexterous manipulation"],
            ["格式", "HDF5 + 多模态同步（RGB、深度、力矩、触觉）"],
            ["许可", "CC-BY-NC 4.0（非商业）"],
            ["访问", "rh20t.github.io"],
        ],
        detail_bullets=[
            "核心差异：除 RGB 外，提供 force-torque 与 tactile 同步流。",
            "适合 contact-rich manipulation：插拔、装配、精细对准。",
            "相机多视角 + 运动信息；标注含 task 与 sub-task。",
        ],
        strengths=[
            "多模态传感在开源 manipulation 数据集中领先。",
            "对研究「接触相位」检测、力控策略价值高。",
        ],
        limitations=[
            "NC 许可限制商业使用。",
            "场景数少于 DROID；跨场景泛化非首要设计目标。",
            "模型需额外 fusion 架构，训练复杂度高于纯视觉 BC。",
        ],
        usage=[
            "目标：contact-aware policy、force-based correction → 首选或补充 DROID。",
            "纯视觉 VLA 若不做力控，RH20T 价值有限。",
            "与 residual RL 结合：力传感可作 reward / observation 扩展参考。",
        ],
    )

    add_dataset_section(
        doc,
        name="2.7 LIBERO（仿真 Benchmark）",
        overview_rows=[
            ["发布方", "LIBERO 团队（NeurIPS 2023 benchmark）"],
            ["规模", "130 tasks；4 个子 benchmark（Spatial/Object/Goal/Long）"],
            ["环境", "MuJoCo / Robosuite 仿真 Franka"],
            ["格式", "HDF5；LIBERO2lerobot / any4lerobot 可转 LeRobot"],
            ["许可", "MIT"],
            ["访问", "libero-project.github.io"],
        ],
        detail_bullets=[
            "不是大规模 pretrain 集，而是标准化 eval benchmark（终身学习、知识迁移）。",
            "衡量：spatial generalization、object generalization、goal generalization、long-horizon。",
            "广泛用于对比 VLA 在仿真中的 sample efficiency 与 transfer。",
        ],
        strengths=[
            "任务设计系统，可重复 eval；无需真机即可对比算法。",
            "与 pickcup 同属仿真验证链路，但任务定义更标准化。",
        ],
        limitations=[
            "Sim-to-real gap 明显；仿真成功 ≠ 真机成功。",
            "视觉域与 MuJoCo pickcup 场景不同，不可直接比数字。",
        ],
        usage=[
            "目标：算法论文 benchmark、architecture ablation → 仿真 eval 标准集。",
            "pickcup 与之互补：pickcup 是 domain-specific G1 抓杯；LIBERO 是 general manipulation benchmark。",
        ],
    )

    add_dataset_section(
        doc,
        name="2.8 ARIO（All Robots In One）",
        overview_rows=[
            ["发布方", "鹏城实验室等（2024）"],
            ["规模", "~300 万轨迹；258 场景；345 skills"],
            ["机器人", "主从双臂等多种"],
            ["格式", "聚合型；部分 LeRobot 转换在 HuggingFace"],
            ["许可", "见官方说明"],
            ["访问", "ario dataset 官方 portal / HuggingFace"],
        ],
        detail_bullets=[
            "中国推出的超大规模聚合数据集，定位类似 OXE 的本土版。",
            "强调多场景、多技能、多机器人统一索引。",
            "生态较新，community tooling 仍在完善。",
        ],
        strengths=["规模大；中文场景与本土 robot 覆盖潜力高。"],
        limitations=["文档与 LeRobot 成熟度不及 OXE/DROID；需自行验证子集质量。"],
        usage=[
            "目标：中国区 generalist pretrain 候选；建议先小规模抽样验证再全量训练。",
            "与 AgiBot 对比：ARIO 偏聚合；AgiBot 偏单组织高质量 factory data。",
        ],
    )

    doc.add_page_break()

    doc.add_heading("3. 次要但常用的数据集", level=1)
    add_table(
        doc,
        ["数据集", "规模", "一句话", "何时用"],
        [
            ["RT-1 Dataset", "130k", "Google 早期 kitchen 数据", "理解 RT 系列历史 baseline"],
            ["RoboSet", "98.5k", "CMU 单臂 + scripted", "多源混合 pretrain 的一小块"],
            ["BC-Z", "26k", "单任务擦桌子", "极窄 task overfitting 测试"],
            ["Dobb-E", "5.6k", "iPhone 采集 home 数据", "低成本 home robot 研究"],
            ["Ego4D", "视频为主", "人类第一视角，无 robot action", "VLA 视觉 pretrain，非 IL 直接"],
            ["NVIDIA GR00T Sim", "仿真", "Isaac 合成 X-Embodiment", "GR00T 官方 sim fine-tune"],
            ["pickcup_train_200", "200 ep", "G1 MuJoCo scripted", "本项目 domain ACT 训练"],
        ],
        col_widths=[3.0, 2.0, 5.5, 5.9],
    )

    doc.add_heading("4. 格式与训练 Pipeline 对照", level=1)
    add_body(
        doc,
        "在 LeRobot / ACT / π0 / GR00T 生态中，「能下载」≠「能直接训」。"
        "关键是 observation.keys、action 语义、fps、normalization stats 与 embodiment tag 一致。",
    )
    add_table(
        doc,
        ["数据集", "原生格式", "LeRobot 路径", "ACT 直接可用", "注意"],
        [
            ["OXE", "RLDS", "IPEC openx-lerobot", "需过滤子集+转换", "action 异构"],
            ["DROID", "RLDS", "lerobot/droid_1.0.1", "是（改 config）", "video key 映射"],
            ["AgiBot", "H5+JSON", "agibot2lerobot", "转换后可用", "v2→v3 API"],
            ["RoboMIND", "自定义", "robomind2lerobot", "部分子集", "EEF vs joint"],
            ["Bridge V2", "TFRecord", "社区脚本", "是", "WidowX 维度"],
            ["pickcup", "LeRobot v3", "原生", "是", "8维 G1 专用"],
        ],
        col_widths=[2.4, 2.2, 3.4, 2.6, 5.8],
    )
    add_caption(doc, "推荐转换工具：github.com/Tavish9/any4lerobot（含 OXE/AgiBot/RoboMIND/LIBERO）")

    doc.add_heading("5. 按目标选型指南", level=1)
    add_table(
        doc,
        ["你的目标", "首选数据集", "备选", "不建议"],
        [
            ["训 cross-robot VLA 基座", "OXE + AgiBot", "ARIO", "仅 Bridge V2"],
            ["单臂真实场景泛化", "DROID", "OXE Franka 子集", "纯仿真"],
            ["人形机长程任务", "AgiBot World", "RoboMIND 天工", "DROID"],
            ["力控 / 接触丰富", "RH20T", "AgiBot（部分 wrench）", "纯 RGB 集"],
            ["语言分层 / 失败学习", "RoboMIND", "AgiBot JSON 标注", "RT-1"],
            ["快速 BC pipeline 验证", "Bridge V2", "pickcup 200", "全量 OXE"],
            ["仿真 benchmark 论文", "LIBERO", "pickcup eval", "DROID 当 sim"],
            ["G1 抓杯本项目", "pickcup_train_200", "—", "混训 OXE 不转换"],
        ],
        col_widths=[4.0, 3.5, 3.5, 4.4],
    )

    doc.add_heading("6. 使用建议（通用原则）", level=1)
    add_bullets(doc, [
        "先明确 embodiment：目标机器人与数据集的 action 空间是否一致或可映射。",
        "Pretrain + Fine-tune 两阶段：大集学表征，小 in-domain 集（如 pickcup 200）学任务。",
        "永远检查 meta/info.json 的 features、fps、stats；混训前各集单独 normalize 再合并。",
        "视频 decode 是 Mac 训练瓶颈；大数据集优先 Linux + CUDA，或预提取帧。",
        "许可证：RH20T NC、部分子集限制商业；产品化前需 legal review。",
        "Eval 与 Train 分离：DROID 场景 hold-out、LIBERO suite 分开报告，避免过拟合数字。",
        "数据质量 > 数量：200 条严格成功 demo 可胜过 2 万条低质 scripted（pickcup 教训）。",
    ])

    doc.add_heading("7. pickcup 项目定位与数据策略建议", level=1)
    add_body(
        doc,
        "pickcup_train_200 是 Unitree G1 23DOF 在 MuJoCo 中的 domain-specific 小集："
        "8 维 proprio + head_camera + 8 维 joint target，LeRobot v3，200 条严格成功 scripted demo。",
        indent=True,
    )
    add_table(
        doc,
        ["阶段", "数据策略", "说明"],
        [
            ["当前", "pickcup 200 + ACT 80k", "跑通 IL 链路与 sim eval"],
            ["短期", "ACT checkpoint eval", "不急于引入 OXE/AgiBot"],
            ["中期", "G1 真机 teleop 50–200 条", "in-domain 微调，质量优先"],
            ["长期", "可选 OXE Franka 子集 pretrain", "仅当扩展到 generalist VLA"],
        ],
        col_widths=[2.5, 4.5, 9.4],
    )
    add_bullets(doc, [
        "不要将 pickcup 与 OXE/AgiBot 无转换混训：action 维度、相机、语义完全不同。",
        "DROID/AgiBot 的价值在于训练 recipe 与 eval 方法论，而非直接提升 G1 抓杯成功率。",
        "若 ACT eval 不佳，优先加 in-domain 数据或 residual RL，而非盲目下载 TB 级数据。",
    ])

    doc.add_heading("8. 结论", level=1)
    add_bullets(doc, [
        "OXE、DROID、AgiBot World 是当前最具影响力的三极：跨机 pretrain、真实场景单臂、人形长程。",
        "RoboMIND、RH20T、Bridge 分别在 multi-emb benchmark、力触觉、BC 入门上有独特价值。",
        "选型由泛化轴决定：embodiment / scene / contact / language / scale。",
        "LeRobot 已成为事实上的训练交换格式，但转换与 embodiment tag 仍是工程主战场。",
        "pickcup 200 条是合格的 domain IL 起点；与 mega-dataset 互补而非竞争。",
    ])

    doc.add_heading("参考资料", level=1)
    add_bullets(doc, [
        "Open X-Embodiment: https://robotics-transformer-x.github.io/",
        "DROID: https://droid-dataset.github.io/",
        "AgiBot World: https://huggingface.co/agibot-world",
        "RoboMIND: https://arxiv.org/html/2412.13877",
        "BridgeData V2: https://rail.eecs.berkeley.edu/datasets/",
        "RH20T: https://rh20t.github.io/",
        "LIBERO: https://libero-project.github.io/",
        "any4lerobot: https://github.com/Tavish9/any4lerobot",
        "LeRobot: https://github.com/huggingface/lerobot",
        "picktask/readme.md",
    ])

    return doc


def main() -> None:
    doc = build_report()
    doc.save(OUTPUT_PATH)
    print(f"已生成报告: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
